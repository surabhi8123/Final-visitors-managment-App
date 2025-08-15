from rest_framework import viewsets, status, filters
from rest_framework.decorators import action
from rest_framework.response import Response
from django_filters.rest_framework import DjangoFilterBackend
from django.db.models import Q
from django.http import HttpResponse
from django.shortcuts import render, redirect
import csv
import io
import base64
from datetime import datetime
from django.conf import settings
from django.utils import timezone
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from .models import Visitor, Visit, VisitorPhoto, CustomAdmin
from .serializers import (
    VisitorSerializer, VisitSerializer, CheckInSerializer, CheckOutSerializer,
    VisitHistorySerializer, VisitorPhotoSerializer
)


class VisitorViewSet(viewsets.ModelViewSet):
    """ViewSet for visitor management."""
    queryset = Visitor.objects.all()
    serializer_class = VisitorSerializer
    filter_backends = [DjangoFilterBackend, filters.SearchFilter, filters.OrderingFilter]
    filterset_fields = ['email', 'phone']
    search_fields = ['name', 'email', 'phone']
    ordering_fields = ['name', 'created_at', 'total_visits']

    @action(detail=False, methods=['post'])
    def check_in(self, request):
        """Check in a new visitor or existing visitor."""
        serializer = CheckInSerializer(data=request.data)
        if serializer.is_valid():
            data = serializer.validated_data
            existing_visitor = data.get('existing_visitor')
            
            # If visitor exists, update their information
            if existing_visitor:
                visitor = existing_visitor
                visitor.name = data['name']
                visitor.email = data['email']
                visitor.phone = data['phone']
                visitor.save()
            else:
                # Create new visitor
                visitor = Visitor.objects.create(
                    name=data['name'],
                    email=data['email'],
                    phone=data['phone']
                )
            
            # Create visit record with signature data if available
            visit_data = {
                'visitor': visitor,
                'purpose': data['purpose'],
                'host_name': data.get('host_name', '')
            }
            
            # Create the visit record first
            visit = Visit.objects.create(**visit_data)
            
            # Handle signature data if provided (either as direct file upload or base64 data)
            signature_file = request.FILES.get('signature_image')
            signature_data = data.get('signature_data')
            is_vector_signature = data.get('is_vector_signature', '').lower() == 'true'
            
            try:
                if signature_file:
                    # Handle file upload
                    print(f"Processing signature file upload: {signature_file.name}")
                    visit.signature_image.save(
                        f"signature_{timezone.now().strftime('%Y%m%d_%H%M%S')}_{uuid.uuid4().hex[:8]}.png",
                        signature_file,
                        save=False
                    )
                    if signature_data:
                        visit.signature_data = signature_data
                    visit.save()
                    print(f"Successfully saved signature file for visit {visit.id}")
                
                elif signature_data:
                    print("Processing signature data...")
                    
                    # Handle vector-based signature
                    if is_vector_signature or (signature_data.startswith('{') and 'paths' in signature_data):
                        print("Detected vector signature data")
                        visit.signature_data = signature_data
                        visit.signature_type = 'vector'
                        visit.save()
                        print(f"Successfully saved vector signature for visit {visit.id}")
                    
                    # Handle image data URL
                    elif signature_data.startswith('data:image/'):
                        print("Processing image signature data...")
                        visit.signature_data = signature_data
                        
                        if ';base64,' in signature_data:
                            format, imgstr = signature_data.split(';base64,')
                            ext = format.split('/')[-1] if '/' in format else 'png'
                            
                            filename = f"signature_{timezone.now().strftime('%Y%m%d_%H%M%S')}_{uuid.uuid4().hex[:8]}.{ext}"
                            
                            try:
                                # Save the signature image file
                                visit.signature_image.save(
                                    filename, 
                                    ContentFile(base64.b64decode(imgstr)), 
                                    save=True
                                )
                                print(f"Successfully saved signature image: {visit.signature_image.name}")
                            except Exception as img_error:
                                print(f"Error saving signature image: {str(img_error)}")
                                print(f"Signature data length: {len(signature_data)}")
                                print(f"Starts with: {signature_data[:50]}...")
                        else:
                            print(f"Invalid signature data format. Expected data URL with base64")
                        
                        visit.save()
                        print(f"Successfully saved visit {visit.id} with image signature")
                    
                    # Handle other types of signature data
                    else:
                        print("Saving raw signature data")
                        visit.signature_data = signature_data
                        visit.save()
                        print(f"Successfully saved raw signature data for visit {visit.id}")
                
                else:
                    print("No signature data or file provided")
                    visit.save()
                    
            except Exception as e:
                import traceback
                error_trace = traceback.format_exc()
                print(f"Error processing signature: {str(e)}\n{error_trace}")
                # Save the visit without the signature data
                visit.save()
            
            # Handle photo upload if provided
            photo_file = request.FILES.get('photo')
            photo_data = data.get('photo_data')
            
            if photo_file or photo_data:
                # Create photo record
                photo_data_dict = {
                    'visitor': visitor,
                    'visit': visit,
                }
                
                if photo_data:
                    # Handle base64 photo data
                    photo_data_dict['image_data'] = photo_data
                elif photo_file:
                    # Handle file upload
                    photo_data_dict['image'] = photo_file
                
                photo = VisitorPhoto.objects.create(**photo_data_dict)
                print(f"Photo saved: {photo.image.url}")
                print(f"Photo path: {photo.image.path}")
                print(f"Photo name: {photo.image.name}")
            
            # Return visit details with photos
            visit_serializer = VisitSerializer(visit, context={'request': request})
            return Response({
                'message': 'Visitor checked in successfully',
                'visit': visit_serializer.data,
                'is_returning_visitor': existing_visitor is not None
            }, status=status.HTTP_201_CREATED)
        
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    @action(detail=False, methods=['post'])
    def check_out(self, request):
        """Check out a visitor."""
        serializer = CheckOutSerializer(data=request.data)
        if serializer.is_valid():
            try:
                visit = Visit.objects.get(id=serializer.validated_data['visit_id'])
                if visit.check_out_time:
                    return Response({
                        'error': 'Visitor has already been checked out'
                    }, status=status.HTTP_400_BAD_REQUEST)
                
                visit.check_out()
                visit_serializer = VisitSerializer(visit)
                return Response({
                    'message': 'Visitor checked out successfully',
                    'visit': visit_serializer.data
                }, status=status.HTTP_200_OK)
            except Visit.DoesNotExist:
                return Response({
                    'error': 'Visit not found'
                }, status=status.HTTP_404_NOT_FOUND)
        
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    @action(detail=False, methods=['get'])
    def active(self, request):
        """Get all currently active visitors."""
        active_visits = Visit.objects.filter(check_out_time__isnull=True)
        serializer = VisitSerializer(active_visits, many=True, context={'request': request})
        return Response({
            'active_visitors': serializer.data,
            'count': active_visits.count()
        })

    @action(detail=False, methods=['get'])
    def history(self, request):
        """Get visit history with search and filter options."""
        visits = Visit.objects.all()
        
        # Apply filters
        name = request.query_params.get('name')
        phone = request.query_params.get('phone')
        email = request.query_params.get('email')
        date_from = request.query_params.get('date_from')
        date_to = request.query_params.get('date_to')
        
        if name:
            visits = visits.filter(visitor__name__icontains=name)
        if phone:
            visits = visits.filter(visitor__phone__icontains=phone)
        if email:
            visits = visits.filter(visitor__email__icontains=email)
        if date_from:
            visits = visits.filter(check_in_time__date__gte=date_from)
        if date_to:
            visits = visits.filter(check_in_time__date__lte=date_to)
        
        # Order by check-in time (newest first)
        visits = visits.order_by('-check_in_time')
        
        serializer = VisitHistorySerializer(visits, many=True, context={'request': request})
        return Response({
            'visits': serializer.data,
            'count': visits.count()
        })

    @action(detail=False, methods=['get'])
    def export(self, request):
        """Export visit history as a Word document (.docx)."""
        visits = Visit.objects.all().order_by('-check_in_time')
        
        # Apply same filters as history endpoint
        name = request.query_params.get('name')
        phone = request.query_params.get('phone')
        email = request.query_params.get('email')
        date_from = request.query_params.get('date_from')
        date_to = request.query_params.get('date_to')
        
        if name:
            visits = visits.filter(visitor__name__icontains=name)
        if phone:
            visits = visits.filter(visitor__phone__icontains=phone)
        if email:
            visits = visits.filter(visitor__email__icontains=email)
        if date_from:
            visits = visits.filter(check_in_time__date__gte=date_from)
        if date_to:
            visits = visits.filter(check_in_time__date__lte=date_to)
        
        # Create a new Word document
        doc = Document()
        
        # Add a title
        title = doc.add_heading('Visit History', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add export date
        export_date = doc.add_paragraph(f'Exported on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        export_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Add a table
        table = doc.add_table(rows=1, cols=17)  # 17 columns as per the original Excel
        table.style = 'Table Grid'
        
        # Set table headers
        headers = [
            'Visitor ID', 'Visitor Name', 'Email', 'Phone Number', 
            'Purpose of Visit', 'Host Name', 'Approval Status', 'Visit ID',
            'Check-in Time', 'Check-out Time', 'Duration (minutes)', 
            'Duration (formatted)', 'Is Active', 'Photo URL', 'Visitor Created At',
            'Visitor Updated At', 'Visit Created At'
        ]
        
        # Style and add headers
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            hdr_cells[i].paragraphs[0].runs[0].bold = True
            hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add data rows
        for visit in visits:
            # Get photo URL if available
            photo_url = ""
            if visit.photos.exists():
                photo = visit.photos.first()
                if photo.image:
                    photo_url = f"http://{request.get_host()}{photo.image.url}"
            
            # Prepare row data
            row_data = [
                str(visit.visitor.id),
                visit.visitor.name,
                visit.visitor.email,
                visit.visitor.phone,
                visit.purpose,
                visit.host_name or 'N/A',
                visit.status,
                str(visit.id),
                visit.check_in_time.strftime('%Y-%m-%d %H:%M:%S'),
                visit.check_out_time.strftime('%Y-%m-%d %H:%M:%S') if visit.check_out_time else 'N/A',
                str(visit.duration_minutes) if visit.duration_minutes is not None else 'N/A',
                visit.duration_formatted,
                'Yes' if visit.is_active else 'No',
                photo_url,
                visit.visitor.created_at.strftime('%Y-%m-%d %H:%M:%S'),
                visit.visitor.updated_at.strftime('%Y-%m-%d %H:%M:%S'),
                visit.check_in_time.strftime('%Y-%m-%d %H:%M:%S')  # Visit created at is same as check-in time
            ]
            
            # Add a new row to the table
            row_cells = table.add_row().cells
            for i, value in enumerate(row_data):
                row_cells[i].text = str(value)
        
        # Auto-fit the table columns
        for section in doc.sections:
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
        
        # Save the document to a bytes buffer
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        
        # Convert to base64 for JSON response
        doc_data = base64.b64encode(doc_buffer.getvalue()).decode('utf-8')
        doc_buffer.close()
        
        # Return JSON response with document data
        return Response({
            'message': 'Visit history exported successfully',
            'filename': f'visit_history_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx',
            'data': doc_data,
            'count': visits.count()
        })

    @action(detail=False, methods=['get'])
    def search(self, request):
        """Search for existing visitors by email or phone."""
        try:
            email = request.query_params.get('email')
            phone = request.query_params.get('phone')
            
            if not email and not phone:
                return Response({
                    'error': 'Please provide email or phone number'
                }, status=status.HTTP_400_BAD_REQUEST)
            
            visitor = None
            if email:
                try:
                    visitor = Visitor.objects.get(email=email)
                except Visitor.DoesNotExist:
                    pass
            
            if not visitor and phone:
                try:
                    visitor = Visitor.objects.get(phone=phone)
                except Visitor.DoesNotExist:
                    pass
            
            if visitor:
                serializer = VisitorSerializer(visitor, context={'request': request})
                return Response({
                    'found': True,
                    'visitor': serializer.data
                })
            else:
                return Response({
                    'found': False,
                    'message': 'No existing visitor found'
                })
        except Exception as e:
            return Response({
                'error': f'Search failed: {str(e)}'
            }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    @action(detail=False, methods=['get'])
    def debug_photos(self, request):
        """Debug endpoint to check photo storage and URLs."""
        try:
            # Get all photos
            photos = VisitorPhoto.objects.all().order_by('-created_at')[:5]
            photo_data = []
            
            for photo in photos:
                photo_data.append({
                    'id': str(photo.id),
                    'visitor_name': photo.visitor.name,
                    'visit_id': str(photo.visit.id),
                    'image_name': photo.image.name,
                    'image_url': photo.image.url,
                    'full_url': request.build_absolute_uri(photo.image.url),
                    'created_at': photo.created_at.isoformat(),
                })
            
            return Response({
                'total_photos': VisitorPhoto.objects.count(),
                'recent_photos': photo_data,
                'media_url': settings.MEDIA_URL,
                'media_root': settings.MEDIA_ROOT,
            })
        except Exception as e:
            return Response({
                'error': str(e)
            }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


class VisitViewSet(viewsets.ModelViewSet):
    """ViewSet for visit management."""
    queryset = Visit.objects.all()
    serializer_class = VisitSerializer
    filter_backends = [DjangoFilterBackend, filters.SearchFilter, filters.OrderingFilter]
    filterset_fields = ['visitor', 'check_in_time', 'check_out_time']
    search_fields = ['visitor__name', 'visitor__email', 'visitor__phone', 'purpose']
    ordering_fields = ['check_in_time', 'check_out_time', 'duration_minutes']


def home(request):
    """Simple home page view."""
    return HttpResponse("<h1>Welcome to ThorSignia visitors management system</h1>")


def admin_login(request):
    """Custom admin login view."""
    if request.method == 'POST':
        email = request.POST.get('email')
        password = request.POST.get('password')
        
        try:
            admin_user = CustomAdmin.objects.get(email=email, is_active=True)
            if admin_user.password == password:
                # Update last login time
                admin_user.last_login = timezone.now()
                admin_user.save()
                
                # Set session
                request.session['admin_logged_in'] = True
                request.session['admin_email'] = email
                
                return redirect('admin_dashboard')
            else:
                return render(request, 'admin/login.html', {
                    'error_message': 'Invalid email or password.'
                })
        except CustomAdmin.DoesNotExist:
            return render(request, 'admin/login.html', {
                'error_message': 'Invalid email or password.'
            })
    
    return render(request, 'admin/login.html')


def admin_dashboard(request):
    """Admin dashboard view."""
    if not request.session.get('admin_logged_in'):
        return redirect('admin_login')
    
    return render(request, 'admin/dashboard.html')


def admin_logout(request):
    """Admin logout view."""
    request.session.flush()
    return redirect('admin_login') 